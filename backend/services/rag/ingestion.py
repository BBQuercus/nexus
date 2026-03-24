"""Document parsing, chunking, and ingestion pipeline."""

from __future__ import annotations

import io
import uuid
from dataclasses import dataclass, field
from typing import Any

import tiktoken

from backend.config import settings
from backend.logging_config import get_logger

logger = get_logger("rag.ingestion")

_enc = tiktoken.get_encoding("cl100k_base")

SUPPORTED_EXTENSIONS = {
    "txt", "md", "csv", "tsv", "xlsx", "xls",
    "pdf", "docx", "pptx", "json",
}

MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB


@dataclass
class ParsedChunk:
    """A single chunk of text extracted from a document."""
    content: str
    chunk_index: int
    page_number: int | None = None
    section_title: str | None = None
    token_count: int = 0
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Full result of parsing a document."""
    filename: str
    content_type: str
    raw_text: str
    chunks: list[ParsedChunk]
    page_count: int | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


def count_tokens(text: str) -> int:
    return len(_enc.encode(text))


def _chunk_text(
    text: str,
    chunk_size: int | None = None,
    chunk_overlap: int | None = None,
) -> list[str]:
    """Split text into overlapping token-bounded chunks."""
    chunk_size = chunk_size or settings.RAG_CHUNK_SIZE
    chunk_overlap = chunk_overlap or settings.RAG_CHUNK_OVERLAP

    tokens = _enc.encode(text)
    if len(tokens) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(tokens):
        end = min(start + chunk_size, len(tokens))
        chunk_tokens = tokens[start:end]
        chunks.append(_enc.decode(chunk_tokens))
        if end >= len(tokens):
            break
        start += chunk_size - chunk_overlap

    return chunks


# ── Parsers by file type ──


def parse_plaintext(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse TXT/MD files — chunk by paragraphs then by token size."""
    text = file_bytes.decode("utf-8", errors="replace")
    total_tokens = count_tokens(text)

    if total_tokens > settings.RAG_MAX_DOCUMENT_TOKENS:
        raise ValueError(
            f"Document exceeds max token limit ({total_tokens:,} > {settings.RAG_MAX_DOCUMENT_TOKENS:,})"
        )

    # Split by double newlines (paragraphs), then merge small ones
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    merged: list[str] = []
    current = ""
    for para in paragraphs:
        candidate = f"{current}\n\n{para}" if current else para
        if count_tokens(candidate) > settings.RAG_CHUNK_SIZE and current:
            merged.append(current)
            current = para
        else:
            current = candidate
    if current:
        merged.append(current)

    # Further split any oversized merged chunks
    chunks: list[ParsedChunk] = []
    idx = 0
    for block in merged:
        for sub in _chunk_text(block):
            chunks.append(ParsedChunk(
                content=sub,
                chunk_index=idx,
                token_count=count_tokens(sub),
            ))
            idx += 1

    return ParsedDocument(
        filename=filename,
        content_type="text/plain",
        raw_text=text,
        chunks=chunks,
        metadata={"total_tokens": total_tokens},
    )


def parse_json(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse JSON files as plain text."""
    return parse_plaintext(file_bytes, filename)


def parse_csv(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse CSV/TSV with schema-aware chunking.

    Produces:
    1. A schema chunk describing columns, types, row count, and summary stats
    2. Row-group chunks with the actual data
    """
    import pandas as pd

    sep = "\t" if filename.endswith((".tsv", ".TSV")) else ","
    df = pd.read_csv(io.BytesIO(file_bytes), sep=sep)

    total_text = df.to_string()
    total_tokens = count_tokens(total_text)
    if total_tokens > settings.RAG_MAX_DOCUMENT_TOKENS:
        raise ValueError(
            f"Document exceeds max token limit ({total_tokens:,} > {settings.RAG_MAX_DOCUMENT_TOKENS:,})"
        )

    chunks: list[ParsedChunk] = []
    idx = 0

    # Schema chunk
    col_info: list[str] = []
    for col in df.columns:
        dtype = str(df[col].dtype)
        nunique = df[col].nunique()
        sample_vals = df[col].dropna().head(3).tolist()
        col_info.append(f"  - {col} ({dtype}, {nunique} unique): samples {sample_vals}")

    schema_text = (
        f'File "{filename}": {len(df)} rows, {len(df.columns)} columns.\n'
        f"Columns:\n" + "\n".join(col_info)
    )

    # Add summary stats for numeric columns
    numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
    if numeric_cols:
        stats = df[numeric_cols].describe().round(2)
        schema_text += f"\n\nSummary statistics:\n{stats.to_string()}"

    chunks.append(ParsedChunk(
        content=schema_text,
        chunk_index=idx,
        token_count=count_tokens(schema_text),
        section_title="Schema & Statistics",
        metadata={"type": "schema"},
    ))
    idx += 1

    # Row-group chunks (50-100 rows)
    rows_per_chunk = 75
    for start_row in range(0, len(df), rows_per_chunk):
        end_row = min(start_row + rows_per_chunk, len(df))
        subset = df.iloc[start_row:end_row]
        row_text = (
            f'Rows {start_row + 1}-{end_row} of "{filename}":\n'
            f"{subset.to_string()}"
        )
        tokens = count_tokens(row_text)
        # If row chunk is too large, further split it
        if tokens > settings.RAG_CHUNK_SIZE * 2:
            for sub in _chunk_text(row_text):
                chunks.append(ParsedChunk(
                    content=sub,
                    chunk_index=idx,
                    token_count=count_tokens(sub),
                    section_title=f"Rows {start_row + 1}-{end_row}",
                    metadata={"type": "data", "row_range": [start_row, end_row]},
                ))
                idx += 1
        else:
            chunks.append(ParsedChunk(
                content=row_text,
                chunk_index=idx,
                token_count=tokens,
                section_title=f"Rows {start_row + 1}-{end_row}",
                metadata={"type": "data", "row_range": [start_row, end_row]},
            ))
            idx += 1

    return ParsedDocument(
        filename=filename,
        content_type="text/csv",
        raw_text=schema_text,  # Store schema as the representative text
        chunks=chunks,
        metadata={
            "rows": len(df),
            "columns": list(df.columns),
            "total_tokens": total_tokens,
        },
    )


def parse_excel(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse Excel files with schema-aware chunking per sheet."""
    import pandas as pd

    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    all_chunks: list[ParsedChunk] = []
    all_schema_parts: list[str] = []
    idx = 0
    total_rows = 0

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        if df.empty:
            continue
        total_rows += len(df)

        # Schema chunk per sheet
        col_info: list[str] = []
        for col in df.columns:
            dtype = str(df[col].dtype)
            nunique = df[col].nunique()
            sample_vals = df[col].dropna().head(3).tolist()
            col_info.append(f"  - {col} ({dtype}, {nunique} unique): samples {sample_vals}")

        schema_text = (
            f'Sheet "{sheet_name}" in "{filename}": {len(df)} rows, {len(df.columns)} columns.\n'
            f"Columns:\n" + "\n".join(col_info)
        )

        numeric_cols = df.select_dtypes(include=["number"]).columns.tolist()
        if numeric_cols:
            stats = df[numeric_cols].describe().round(2)
            schema_text += f"\n\nSummary statistics:\n{stats.to_string()}"

        all_schema_parts.append(schema_text)
        all_chunks.append(ParsedChunk(
            content=schema_text,
            chunk_index=idx,
            token_count=count_tokens(schema_text),
            section_title=f'Sheet "{sheet_name}" — Schema',
            metadata={"type": "schema", "sheet": sheet_name},
        ))
        idx += 1

        # Row-group chunks
        rows_per_chunk = 75
        for start_row in range(0, len(df), rows_per_chunk):
            end_row = min(start_row + rows_per_chunk, len(df))
            subset = df.iloc[start_row:end_row]
            row_text = (
                f'Sheet "{sheet_name}", Rows {start_row + 1}-{end_row}:\n'
                f"{subset.to_string()}"
            )
            for sub in _chunk_text(row_text):
                all_chunks.append(ParsedChunk(
                    content=sub,
                    chunk_index=idx,
                    token_count=count_tokens(sub),
                    section_title=f'Sheet "{sheet_name}" — Rows {start_row + 1}-{end_row}',
                    metadata={"type": "data", "sheet": sheet_name, "row_range": [start_row, end_row]},
                ))
                idx += 1

    full_schema = "\n\n".join(all_schema_parts)
    total_tokens = sum(c.token_count for c in all_chunks)
    if total_tokens > settings.RAG_MAX_DOCUMENT_TOKENS:
        raise ValueError(
            f"Document exceeds max token limit ({total_tokens:,} > {settings.RAG_MAX_DOCUMENT_TOKENS:,})"
        )

    return ParsedDocument(
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        raw_text=full_schema,
        chunks=all_chunks,
        page_count=len(xls.sheet_names),
        metadata={
            "sheets": xls.sheet_names,
            "total_rows": total_rows,
            "total_tokens": total_tokens,
        },
    )


def parse_pdf(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse PDF files.

    Attempts Docling first for high-quality table extraction.
    Falls back to a simple page-by-page text extraction.
    """
    try:
        return _parse_pdf_docling(file_bytes, filename)
    except ImportError:
        logger.warning("docling_not_installed_fallback", filename=filename)
        return _parse_pdf_fallback(file_bytes, filename)
    except Exception:
        logger.exception("docling_failed_fallback", filename=filename)
        return _parse_pdf_fallback(file_bytes, filename)


def _parse_pdf_docling(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse PDF using Docling for high-accuracy table extraction."""
    from docling.document_converter import DocumentConverter
    import tempfile
    import os

    # Docling needs a file path
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        doc = result.document

        full_text = doc.export_to_markdown()
        total_tokens = count_tokens(full_text)
        if total_tokens > settings.RAG_MAX_DOCUMENT_TOKENS:
            raise ValueError(
                f"Document exceeds max token limit ({total_tokens:,} > {settings.RAG_MAX_DOCUMENT_TOKENS:,})"
            )

        # Chunk by document structure elements
        chunks: list[ParsedChunk] = []
        idx = 0
        current_section = ""
        current_text = ""

        for item in doc.iterate_items():
            element = item[1] if isinstance(item, tuple) else item
            text = element.export_to_markdown() if hasattr(element, 'export_to_markdown') else str(element)
            if not text.strip():
                continue

            # Detect section headings
            el_type = type(element).__name__
            if el_type in ("SectionHeaderItem", "HeadingItem"):
                # Flush current section
                if current_text.strip():
                    for sub in _chunk_text(current_text):
                        chunks.append(ParsedChunk(
                            content=sub,
                            chunk_index=idx,
                            token_count=count_tokens(sub),
                            section_title=current_section or None,
                            metadata={"type": "section"},
                        ))
                        idx += 1
                current_section = text.strip().lstrip("# ")
                current_text = ""
            else:
                current_text += text + "\n\n"

        # Flush last section
        if current_text.strip():
            for sub in _chunk_text(current_text):
                chunks.append(ParsedChunk(
                    content=sub,
                    chunk_index=idx,
                    token_count=count_tokens(sub),
                    section_title=current_section or None,
                    metadata={"type": "section"},
                ))
                idx += 1

        # If no structured chunks produced, fall back to simple chunking
        if not chunks:
            for sub in _chunk_text(full_text):
                chunks.append(ParsedChunk(
                    content=sub,
                    chunk_index=idx,
                    token_count=count_tokens(sub),
                ))
                idx += 1

        return ParsedDocument(
            filename=filename,
            content_type="application/pdf",
            raw_text=full_text,
            chunks=chunks,
            metadata={"total_tokens": total_tokens, "parser": "docling"},
        )
    finally:
        os.unlink(tmp_path)


def _parse_pdf_fallback(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Simple PDF parsing using pypdf as fallback."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Neither docling nor pypdf is installed for PDF parsing")

    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)

    full_text = "\n\n".join(pages_text)
    total_tokens = count_tokens(full_text)
    if total_tokens > settings.RAG_MAX_DOCUMENT_TOKENS:
        raise ValueError(
            f"Document exceeds max token limit ({total_tokens:,} > {settings.RAG_MAX_DOCUMENT_TOKENS:,})"
        )

    chunks: list[ParsedChunk] = []
    idx = 0
    for page_num, page_text in enumerate(pages_text, 1):
        if not page_text.strip():
            continue
        for sub in _chunk_text(page_text):
            chunks.append(ParsedChunk(
                content=sub,
                chunk_index=idx,
                page_number=page_num,
                token_count=count_tokens(sub),
                metadata={"type": "page"},
            ))
            idx += 1

    return ParsedDocument(
        filename=filename,
        content_type="application/pdf",
        raw_text=full_text,
        chunks=chunks,
        page_count=len(reader.pages),
        metadata={"total_tokens": total_tokens, "parser": "pypdf"},
    )


def parse_docx(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse DOCX files. Try Docling first, fall back to python-docx."""
    try:
        from docling.document_converter import DocumentConverter
        import tempfile, os

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            converter = DocumentConverter()
            result = converter.convert(tmp_path)
            text = result.document.export_to_markdown()
        finally:
            os.unlink(tmp_path)
    except (ImportError, Exception):
        # Fallback: simple paragraph extraction
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("Neither docling nor python-docx is installed")
        doc = DocxDocument(io.BytesIO(file_bytes))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    total_tokens = count_tokens(text)
    if total_tokens > settings.RAG_MAX_DOCUMENT_TOKENS:
        raise ValueError(
            f"Document exceeds max token limit ({total_tokens:,} > {settings.RAG_MAX_DOCUMENT_TOKENS:,})"
        )

    chunks: list[ParsedChunk] = []
    idx = 0
    for sub in _chunk_text(text):
        chunks.append(ParsedChunk(
            content=sub,
            chunk_index=idx,
            token_count=count_tokens(sub),
        ))
        idx += 1

    return ParsedDocument(
        filename=filename,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        raw_text=text,
        chunks=chunks,
        metadata={"total_tokens": total_tokens},
    )


# ── Main entry point ──


PARSERS = {
    "txt": parse_plaintext,
    "md": parse_plaintext,
    "csv": parse_csv,
    "tsv": parse_csv,
    "json": parse_json,
    "xlsx": parse_excel,
    "xls": parse_excel,
    "pdf": parse_pdf,
    "docx": parse_docx,
    "pptx": parse_docx,  # Docling handles PPTX too
}


def parse_document(file_bytes: bytes, filename: str) -> ParsedDocument:
    """Parse a document into chunks based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in PARSERS:
        raise ValueError(f"Unsupported file type: .{ext}")
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(f"File exceeds {MAX_FILE_SIZE // (1024 * 1024)}MB limit")

    logger.info("parsing_document", filename=filename, extension=ext, size_bytes=len(file_bytes))
    return PARSERS[ext](file_bytes, filename)
